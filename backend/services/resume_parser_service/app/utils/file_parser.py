import io
from pathlib import Path

import pytesseract
from docx import Document
from PIL import Image
from pypdf import PdfReader

from app.core.config import settings
from common.logging import get_logger

logger = get_logger(__name__)


class FileParseError(Exception):
    pass


def parse_pdf(file_path: str | Path) -> str:

    try:
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileParseError(f"File not found: {file_path}")

        logger.info(f"Parsing PDF: {file_path}")

        reader = PdfReader(str(file_path))
        text = ""

        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
                logger.debug(f"Extracted {len(page_text)} chars from page {page_num}")

        text = text.strip()

        if not text:
            logger.warning(f"No text extracted from PDF: {file_path}")
            return ""

        logger.info(f"Successfully parsed PDF: {len(text)} characters")
        return text

    except Exception as e:
        logger.error(f"Error parsing PDF {file_path}: {str(e)}")
        raise FileParseError(f"Failed to parse PDF: {str(e)}")


def parse_docx(file_path: str | Path) -> str:

    try:
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileParseError(f"File not found: {file_path}")

        logger.info(f"Parsing DOCX: {file_path}")

        doc = Document(str(file_path))
        text = ""

        for para_num, paragraph in enumerate(doc.paragraphs, 1):
            if paragraph.text.strip():
                text += paragraph.text + "\n"
                logger.debug(f"Extracted paragraph {para_num}")

        # Also extract text from tables
        for table_num, table in enumerate(doc.tables, 1):
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
            text += "\n"
            logger.debug(f"Extracted table {table_num}")

        text = text.strip()

        if not text:
            logger.warning(f"No text extracted from DOCX: {file_path}")
            return ""

        logger.info(f"Successfully parsed DOCX: {len(text)} characters")
        return text

    except Exception as e:
        logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
        raise FileParseError(f"Failed to parse DOCX: {str(e)}")


def parse_with_ocr(file_path: str | Path, language: str = "eng") -> str:

    try:
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileParseError(f"File not found: {file_path}")

        logger.info(f"Parsing with OCR: {file_path}")

        # Configure tesseract
        if settings.tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

        # Open image
        image = Image.open(str(file_path))

        # Perform OCR
        text = pytesseract.image_to_string(image, lang=language)
        text = text.strip()

        if not text:
            logger.warning(f"No text extracted via OCR: {file_path}")
            return ""

        logger.info(f"Successfully parsed with OCR: {len(text)} characters")
        return text

    except Exception as e:
        logger.error(f"Error parsing with OCR {file_path}: {str(e)}")
        raise FileParseError(f"Failed to parse with OCR: {str(e)}")


def parse_file(file_path: str | Path, use_ocr: bool = False) -> str:

    # Convert relative path to absolute path
    file_path = Path(file_path)

    # If path is relative, prepend the upload directory
    if not file_path.is_absolute():
        base_dir = Path(__file__).resolve().parent.parent.parent
        file_path = base_dir / "uploads" / file_path

    extension = file_path.suffix.lower()

    logger.info(f"Parsing file: {file_path} (extension: {extension})")

    try:
        if extension == ".pdf":
            text = parse_pdf(file_path)

            # If no text extracted and OCR is enabled, try OCR
            if not text and use_ocr:
                logger.info("No text in PDF, attempting OCR")
                text = parse_with_ocr(file_path)

            return text

        elif extension in [".docx", ".doc"]:
            return parse_docx(file_path)

        elif extension in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            return parse_with_ocr(file_path)

        else:
            raise FileParseError(f"Unsupported file type: {extension}")

    except FileParseError:
        raise
    except Exception as e:
        logger.error(f"Unexpected error parsing {file_path}: {str(e)}")
        raise FileParseError(f"Failed to parse file: {str(e)}")


def extract_text_from_bytes(
    file_bytes: bytes, filename: str, use_ocr: bool = False
) -> str:

    extension = Path(filename).suffix.lower()

    logger.info(f"Parsing bytes for: {filename} (extension: {extension})")

    try:
        if extension == ".pdf":
            reader = PdfReader(io.BytesIO(file_bytes))
            text = ""

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            return text.strip()

        elif extension in [".docx", ".doc"]:
            doc = Document(io.BytesIO(file_bytes))
            text = ""

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            return text.strip()

        elif extension in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            image = Image.open(io.BytesIO(file_bytes))

            if settings.tesseract_cmd:
                pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

            text = pytesseract.image_to_string(image)
            return text.strip()

        else:
            raise FileParseError(f"Unsupported file type: {extension}")

    except Exception as e:
        logger.error(f"Error parsing bytes for {filename}: {str(e)}")
        raise FileParseError(f"Failed to parse file bytes: {str(e)}")
